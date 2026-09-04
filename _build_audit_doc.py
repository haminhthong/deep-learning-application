from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path("docs/Huong_dan_cai_thien_du_an_PPE_4_tang.docx")
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(89, 89, 89)
RED = RGBColor(192, 0, 0)
GREEN = RGBColor(0, 128, 64)
LIGHT_BLUE = "E8EEF5"


def font(run, size=11, bold=False, color=None, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
                element = OxmlElement(f"w:{side}")
                element.set(qn("w:w"), str(value))
                element.set(qn("w:type"), "dxa")
                margins.append(element)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    header_properties.append(marker)
    for index, title in enumerate(headers):
        shade(table.rows[0].cells[index], LIGHT_BLUE)
        font(table.rows[0].cells[index].paragraphs[0].add_run(title), 9.5, True, DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            font(cells[index].paragraphs[0].add_run(str(value)), 9.5)
    table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        font(paragraph.add_run(bold_prefix), 11, True)
        font(paragraph.add_run(text[len(bold_prefix) :]), 11)
    else:
        font(paragraph.add_run(text), 11)
    return paragraph


def bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    font(paragraph.add_run(text), 11)


def numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    font(paragraph.add_run(text), 11)


def code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(3)
    font(paragraph.add_run(text), 9.5, name="Consolas")


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(paragraph.add_run("Trang "), 9, color=GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    font(header.add_run("PPE SAFETY SURVEILLANCE | LỘ TRÌNH CẢI THIỆN 4 TẦNG"), 8.5, True, GRAY)
    page_number(section.footer.paragraphs[0])


def build():
    doc = Document()
    configure(doc)

    doc.add_paragraph().paragraph_format.space_after = Pt(85)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(kicker.add_run("HƯỚNG DẪN CẢI THIỆN DỰ ÁN COMPUTER VISION"), 10, True, BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(9)
    font(
        title.add_run("Đánh giá và nâng cấp hệ thống\ngiám sát PPE theo 4 tầng"),
        27,
        True,
        DARK_BLUE,
    )
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(
        subtitle.add_run(
            "Problem → AI/ML Correctness → Software Engineering → Production & Business Value"
        ),
        12.5,
        color=GRAY,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(75)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(
        note.add_run(
            "Checklist kỹ thuật • Lộ trình triển khai • Tiêu chí nghiệm thu • Hướng dẫn CV"
        ),
        10.5,
        True,
        DARK_BLUE,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(90)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(
        meta.add_run("Dự án: PPE Safety Surveillance bằng YOLO\nPhiên bản: 1.0 | Năm 2026"),
        10,
        color=GRAY,
    )
    doc.add_page_break()

    heading(doc, "Tóm tắt điều hành", 1)
    body(
        doc,
        "Dự án hiện là một portfolio prototype có cấu trúc phần mềm khá tốt: pipeline hai model, IoU Tracker, xác nhận vi phạm theo thời gian, CLI, Streamlit, báo cáo và kiểm thử. Điểm yếu lớn nhất không nằm ở giao diện mà ở việc chưa có bằng chứng về dataset, training, leakage, baseline và metric thực tế.",
    )
    add_table(
        doc,
        ["Tầng", "Hiện trạng", "Ưu tiên"],
        [
            ("Problem", "Use case hợp lý nhưng thiếu KPI định lượng", "P1"),
            ("AI/ML correctness", "Inference hợp lý; chưa có training evidence", "P0"),
            ("Software Engineering", "Tách module tốt; còn portability/edge cases", "P1"),
            ("Production/Business", "Demo local; chưa chịu tải và bảo mật", "P2"),
        ],
        [2100, 5160, 2100],
    )
    heading(doc, "Nguyên tắc cải thiện", 2)
    for text in (
        "Không thêm tính năng chỉ để project trông phức tạp.",
        "Không công bố metric chưa đo hoặc kết quả từ demo mô phỏng.",
        "Mọi thay đổi phải có tiêu chí nghiệm thu và bằng chứng tái lập.",
        "Ưu tiên dataset, baseline và metric trước dashboard hoặc microservice.",
    ):
        bullet(doc, text)
    doc.add_page_break()

    heading(doc, "1. Tầng Problem - Xác định đúng bài toán", 1)
    heading(doc, "1.1 Problem statement đề xuất", 2)
    body(
        doc,
        "Xây dựng hệ thống phát hiện người không đội mũ bảo hộ hoặc không mặc áo phản quang trong ảnh/video công trường, duy trì ID tạm thời, xác nhận vi phạm qua nhiều lần inference và lưu bằng chứng để hỗ trợ giám sát viên.",
    )
    heading(doc, "1.2 Các câu hỏi bắt buộc phải trả lời", 2)
    for text in (
        "Ai là đối tượng cần giám sát: công nhân, khách hay tất cả người trong ROI?",
        "Camera cố định, webcam hay RTSP; khoảng cách và góc đặt dự kiến?",
        "Mức false alarm và missed violation có thể chấp nhận?",
        "Cảnh báo phải xuất hiện trong bao lâu?",
        "Snapshot được lưu bao lâu và ai được phép xem?",
        "Điều kiện nào được xem là ngoài phạm vi: che khuất, mũ cầm tay, người quá nhỏ?",
    ):
        bullet(doc, text)
    heading(doc, "1.3 KPI mục tiêu", 2)
    add_table(
        doc,
        ["KPI", "Mục tiêu thiết kế", "Cách đo"],
        [
            ("Recall no-helmet", ">= 90%", "Test set độc lập"),
            ("Recall no-vest", ">= 85%", "Test set độc lập"),
            ("Precision cảnh báo", ">= 85%", "Đánh giá theo sự kiện"),
            ("False alerts", "<= 2/giờ/camera", "Video pilot có ground truth"),
            ("Time-to-alert", "<= 2 giây", "Timestamp vi phạm đến alert"),
            ("GPU throughput", ">= 20 FPS", "Đúng phần cứng và image size"),
        ],
        [2400, 2160, 4800],
    )
    body(
        doc,
        "Lưu ý: các số trên là mục tiêu định hướng, không phải kết quả của project cho đến khi có phép đo thực tế.",
    )
    heading(doc, "1.4 Business rules cần đưa vào repository", 2)
    for text in (
        "Chỉ kiểm tra người có tâm bounding box nằm trong ROI.",
        "Không phát hiện helmet không tự động đồng nghĩa no-helmet.",
        "Mỗi cặp track ID và loại vi phạm chỉ tạo một event trong phiên.",
        "Vi phạm video phải xuất hiện đủ N lần inference liên tiếp.",
        "Track bị đổi ID có thể làm phát sinh sự kiện mới và phải được ghi là limitation.",
    ):
        numbered(doc, text)
    doc.add_page_break()

    heading(doc, "2. Tầng AI/ML Correctness", 1)
    heading(doc, "2.1 Cấu trúc training cần bổ sung", 2)
    for line in (
        "training/",
        "├── data.yaml",
        "├── train.py",
        "├── evaluate.py",
        "├── export.py",
        "├── dataset_card.md",
        "├── split_manifest.csv",
        "└── configs/yolov8n_ppe.yaml",
    ):
        code(doc, line)
    heading(doc, "2.2 Dataset card", 2)
    for text in (
        "Nguồn và license dữ liệu.",
        "Số ảnh, video, camera, địa điểm và độ phân giải.",
        "Phân bố instance theo class và split.",
        "Quy trình annotation và xử lý nhãn mơ hồ.",
        "Điều kiện ánh sáng, che khuất và khoảng cách.",
        "Bias, limitation và quyền sử dụng hình ảnh con người.",
    ):
        bullet(doc, text)
    heading(doc, "2.3 Ngăn data leakage", 2)
    body(
        doc,
        "Không chia ngẫu nhiên từng frame. Frame liên tiếp của cùng video gần như giống nhau và làm metric test cao giả tạo. Hãy chia theo nhóm camera_id + location_id + recording_session; một phiên quay chỉ thuộc một split.",
    )
    add_table(
        doc,
        ["Kiểm tra", "Yêu cầu"],
        [
            ("File trùng hoàn toàn", "So sánh SHA-256 giữa các split"),
            ("Ảnh gần giống", "Dùng perceptual hash"),
            ("Frame cùng video", "Không được nằm ở hai split"),
            ("Test set", "Không dùng chọn threshold/hyperparameter"),
            ("Camera test", "Ưu tiên camera/bối cảnh chưa thấy khi train"),
        ],
        [2700, 6660],
    )
    heading(doc, "2.4 Đồng bộ preprocessing", 2)
    body(
        doc,
        "Model PPE trong inference nhận person crop có padding. Training data phải có cùng kiểu crop, padding, image size, class order và normalization. Nếu training dùng ảnh toàn cảnh nhưng inference dùng crop, hệ thống có train-serve skew.",
    )
    for text in (
        "Đưa person_roi_padding vào DetectionConfig.",
        "Dùng cùng giá trị trong script tạo crop training và detector inference.",
        "Lưu phiên bản Ultralytics, PyTorch, seed và image size cho mỗi experiment.",
        "Kiểm tra class ID trong weight trước khi chạy pipeline.",
    ):
        bullet(doc, text)
    doc.add_page_break()

    heading(doc, "2.5 Baseline và ablation study", 2)
    add_table(
        doc,
        ["ID", "Phiên bản", "Mục đích"],
        [
            ("B0", "Một YOLO phát hiện trực tiếp toàn bộ class", "Baseline đơn giản"),
            ("B1", "Two-stage, không tracking", "Đo lợi ích crop người"),
            ("B2", "Two-stage + IoU Tracker", "Đo lợi ích ID association"),
            ("B3", "B2 + temporal confirmation", "Đo giảm false alert"),
            ("B4", "B3 + ByteTrack/BoT-SORT", "Đo cải thiện tracking"),
        ],
        [900, 4800, 3660],
    )
    heading(doc, "2.6 Metric phù hợp", 2)
    add_table(
        doc,
        ["Cấp", "Metric", "Ý nghĩa"],
        [
            ("Detection", "Precision, Recall, F1, AP, mAP", "Chất lượng bounding box/class"),
            ("Alert", "Event Precision/Recall", "Cảnh báo đúng theo sự kiện"),
            ("Alert", "False alerts/hour", "Chi phí vận hành thực tế"),
            ("Latency", "Time-to-alert", "Độ trễ do confirmation"),
            ("Tracking", "ID switches, IDF1", "Ổn định định danh"),
            ("Runtime", "FPS, latency, VRAM", "Khả năng triển khai"),
        ],
        [1500, 3000, 4860],
    )
    heading(doc, "2.7 Logic nhãn xung đột", 2)
    body(
        doc,
        "Nếu cùng ROI xuất hiện helmet và no-helmet, không nên chỉ dùng phép có/không. Nên so confidence và dùng conflict margin được chọn trên validation set.",
    )
    code(doc, "violation = no_helmet_score >= threshold")
    code(doc, "            and no_helmet_score > helmet_score + conflict_margin")
    heading(doc, "2.8 Edge cases cần test", 2)
    for text in (
        "Frame None/rỗng; ảnh hỏng; video FPS bằng 0.",
        "Không có người; người sát biên; bounding box âm hoặc đảo tọa độ.",
        "Hai người overlap cao; hai detection có cùng IoU.",
        "Track mất rồi xuất hiện; violation xuất hiện-mất-xuất hiện.",
        "Helmet và no-helmet cùng xuất hiện; class lạ; tên class khác format.",
        "ROI polygon không hợp lệ; output không có quyền ghi; VideoWriter lỗi.",
        "CUDA out-of-memory và model load failure.",
    ):
        bullet(doc, text)
    doc.add_page_break()

    heading(doc, "3. Tầng Software Engineering", 1)
    heading(doc, "3.1 Sửa portability ngay", 2)
    body(
        doc,
        "pytest.ini đang chứa đường dẫn temp theo tài khoản Windows cá nhân. Hãy bỏ đường dẫn tuyệt đối để người khác clone repository vẫn chạy được.",
    )
    code(doc, "[pytest]")
    code(doc, "testpaths = tests")
    code(doc, "python_files = test_*.py")
    code(doc, "addopts = -v")
    heading(doc, "3.2 Tách dependency", 2)
    body(
        doc,
        "requirements.txt chỉ chứa runtime; requirements-dev.txt chứa pytest, coverage, Ruff và mypy. Việc này giảm dependency khi deploy.",
    )
    heading(doc, "3.3 Quality gates", 2)
    for command in (
        "ruff check .",
        "ruff format --check .",
        "pytest --cov=ppe_detection --cov-report=term-missing",
        "python app.py --help",
        "python app.py --demo --source sample.jpg --save --no-display",
    ):
        code(doc, command)
    heading(doc, "3.4 GitHub Actions", 2)
    for text in (
        "Checkout source và setup Python 3.11.",
        "Cài requirements-dev.txt.",
        "Chạy Ruff và pytest trên push/pull request.",
        "Không tải model thật trong CI; pipeline test dùng SyntheticDemoDetector.",
    ):
        numbered(doc, text)
    heading(doc, "3.5 Làm rõ demo", 2)
    body(
        doc,
        "Không tự động chuyển sang demo khi thiếu model vì người dùng có thể hiểu nhầm kết quả mô phỏng là inference. Chỉ bật demo khi có --demo và hiển thị thông báo nổi bật: PIPELINE SIMULATION - KHÔNG PHẢI KẾT QUẢ AI.",
    )
    heading(doc, "3.6 Upload và file tạm", 2)
    for text in (
        "Giới hạn kích thước và thời lượng video.",
        "Kiểm tra MIME/file signature thay vì chỉ extension.",
        "Xóa file tạm trong finally kể cả khi có exception ngoài dự kiến.",
        "Tạo output directory riêng theo session UUID.",
        "Không để snapshot nằm trong thư mục web public.",
    ):
        bullet(doc, text)
    doc.add_page_break()

    heading(doc, "3.7 Cấu trúc đề xuất", 2)
    for line in (
        "app.py / web_app.py",
        "        ↓",
        "DetectionService",
        "        ↓",
        "ModelBundle (dùng chung, stateless)",
        "SessionPipeline (tracker/report riêng)",
        "        ↓",
        "OutputStore",
    ):
        code(doc, line)
    body(
        doc,
        "Tách model dùng chung khỏi state theo session giúp giảm thời gian load nhưng tránh dùng chung tracker giữa hai camera hoặc hai người dùng.",
    )
    heading(doc, "3.8 Checklist người khác clone repo", 2)
    for text in (
        "Không còn path máy cá nhân.",
        "README có Python version và lệnh cài đặt.",
        "Có sample input nhỏ hoặc link tải.",
        "Có link/checksum model hoặc hướng dẫn train.",
        "Demo ghi rõ là mô phỏng.",
        "pytest và CLI help chạy trên Windows/Linux.",
        "Không commit model, dataset, output hoặc cache.",
    ):
        bullet(doc, text)

    heading(doc, "4. Tầng Production và Business Value", 1)
    heading(doc, "4.1 Vì sao hiện tại chưa chịu được 100 users", 2)
    for text in (
        "Mỗi lượt Streamlit có thể tạo pipeline và load model riêng.",
        "Không có queue, timeout, concurrency limit hay GPU scheduler.",
        "Video inference là tác vụ dài và chặn worker.",
        "Nhiều session có thể ghi chung output directory.",
        "Không có rate limiting, authentication hoặc quota.",
    ):
        bullet(doc, text)
    heading(doc, "4.2 Kiến trúc production đề xuất", 2)
    for line in (
        "Web/Client",
        "    ↓",
        "FastAPI + Authentication + Rate Limit",
        "    ↓",
        "Redis/Celery Job Queue",
        "    ↓",
        "GPU Inference Workers (load model một lần)",
        "    ↓",
        "PostgreSQL + Object Storage",
        "    ↓",
        "Dashboard / Notification Service",
    ):
        code(doc, line)
    heading(doc, "4.3 Camera realtime", 2)
    body(
        doc,
        "Tổ chức worker theo camera, không theo người xem dashboard. Một camera có một inference worker và nhiều dashboard chỉ đọc event stream. Cách này khiến 100 người xem không tạo 100 bản model.",
    )
    heading(doc, "4.4 Security và privacy", 2)
    add_table(
        doc,
        ["Rủi ro", "Biện pháp tối thiểu"],
        [
            ("Upload độc hại/quá lớn", "MIME validation, size limit, timeout"),
            ("Snapshot con người", "Encryption, access control, retention"),
            ("Truy cập trái phép", "Authentication và role-based access"),
            ("Lộ dữ liệu", "HTTPS, private object storage, audit log"),
            ("Disk exhaustion", "Quota, cleanup job, monitoring"),
            ("Privacy", "Consent/signage, blur mặt khi phù hợp"),
        ],
        [3000, 6360],
    )
    doc.add_page_break()

    heading(doc, "4.5 Pilot chứng minh giá trị", 2)
    body(
        doc,
        "Trước khi xây microservice phức tạp, nên chạy pilot nhỏ tại một địa điểm với 2-3 camera trong khoảng hai tuần và có giám sát viên xác nhận từng cảnh báo.",
    )
    for text in (
        "Số giờ video đã xử lý.",
        "Số vi phạm thật và số cảnh báo.",
        "True positive, false positive, false negative.",
        "False alerts/hour và time-to-alert.",
        "Thời gian giám sát thủ công tiết kiệm.",
        "Phản hồi của người vận hành và chi phí hạ tầng.",
    ):
        bullet(doc, text)
    heading(doc, "4.6 Monitoring", 2)
    for text in (
        "FPS và inference latency theo camera.",
        "GPU utilization, VRAM và queue length.",
        "Camera offline và job timeout.",
        "Disk usage và lỗi ghi video/snapshot.",
        "False alert do người vận hành đánh dấu.",
    ):
        bullet(doc, text)

    heading(doc, "5. Lộ trình triển khai", 1)
    add_table(
        doc,
        ["Giai đoạn", "Thời lượng", "Đầu ra bắt buộc"],
        [
            ("1. Repository", "1-2 ngày", "Portability, Ruff, CI, upload safety, session output"),
            ("2. ML evidence", "3-7 ngày", "Dataset card, split, training, baseline, metrics"),
            ("3. Use case", "2-4 ngày", "Video thật, false alerts/hour, latency, examples"),
            ("4. Production prototype", "1-2 tuần", "API, queue, worker, auth, Docker, monitoring"),
        ],
        [2100, 1800, 5460],
    )
    heading(doc, "5.1 Definition of Done", 2)
    for text in (
        "Problem: KPI và business rules được ghi rõ.",
        "AI/ML: test set độc lập, không leakage, có baseline và metric thật.",
        "Software: CI xanh, không path cá nhân, test edge cases quan trọng.",
        "Demo: có video thực, ví dụ đúng/sai và phần cứng benchmark.",
        "Production: có threat/privacy checklist và kiến trúc chịu tải hợp lý.",
        "CV: mọi con số đều truy ngược được về experiment/report.",
    ):
        bullet(doc, text)

    heading(doc, "6. Thứ tự ưu tiên khi đưa vào CV", 1)
    for text in (
        "Dataset card và cách split không leakage.",
        "Metric thật và baseline.",
        "Video demo thực tế.",
        "Test, CI và README tái lập.",
        "Temporal confirmation và reporting.",
        "Streamlit/Docker/API sau khi model đã được chứng minh.",
    ):
        numbered(doc, text)
    heading(doc, "6.1 Mô tả CV mẫu", 2)
    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Inches(0.35)
    quote.paragraph_format.right_indent = Inches(0.35)
    quote.paragraph_format.space_before = Pt(8)
    quote.paragraph_format.space_after = Pt(10)
    font(
        quote.add_run(
            "Xây dựng hệ thống giám sát PPE hai giai đoạn bằng Ultralytics YOLO và OpenCV; phát hiện người, phân tích mũ/áo trên từng ROI, duy trì ID bằng IoU Tracker, xác nhận vi phạm nhiều frame, lưu ảnh bằng chứng và xuất báo cáo JSON/CSV; hỗ trợ CLI, Streamlit và chế độ headless."
        ),
        11,
        italic=True,
        color=DARK_BLUE,
    )
    body(
        doc,
        "Chỉ bổ sung các con số mAP, Recall hoặc FPS sau khi đã đo trên test set và phần cứng được mô tả rõ.",
    )
    heading(doc, "6.2 Câu hỏi phỏng vấn cần chuẩn bị", 2)
    for text in (
        "Vì sao chọn two-stage thay vì một detector duy nhất?",
        "Bạn ngăn leakage giữa các frame video như thế nào?",
        "Temporal confirmation ảnh hưởng Precision và latency ra sao?",
        "IoU Tracker thất bại khi nào và vì sao ByteTrack tốt hơn?",
        "Metric detection khác metric event-level như thế nào?",
        "Nếu có 100 users hoặc nhiều camera, bạn thay đổi kiến trúc ra sao?",
        "Bạn xử lý privacy của snapshot người vi phạm thế nào?",
    ):
        bullet(doc, text)

    heading(doc, "Kết luận", 1)
    body(
        doc,
        "Hướng cải thiện tốt nhất không phải làm project trông nhiều AI hơn, mà là tăng mức độ chứng minh: dữ liệu đúng, split đúng, baseline rõ, metric tái lập, demo thật và giới hạn trung thực. Khi bốn tầng đều có bằng chứng, project sẽ mạnh hơn đáng kể trong CV và phỏng vấn kỹ thuật.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Hướng dẫn cải thiện dự án PPE theo 4 tầng"
    doc.core_properties.subject = (
        "Problem, AI/ML, Software Engineering, Production và Business Value"
    )
    doc.core_properties.author = "Hà Minh Thông"
    doc.save(OUTPUT)
    print("DOCX_CREATED")


if __name__ == "__main__":
    build()
